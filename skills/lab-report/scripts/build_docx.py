"""Helper library for building Chinese/English lab-report .docx files.

Reusable so each report calls the same functions instead of re-writing
python-docx boilerplate. Set the East Asian font explicitly: python-docx
only sets the Latin font by default.

Usage:
    import sys, os
    sys.path.insert(0, os.path.join(SKILL_DIR, "scripts"))
    from build_docx import new_report, add_code, add_table, export_pdf

    doc = new_report()
    doc.add_heading("实验题目", level=1)
    doc.add_paragraph("...")
    add_code(doc, "print('hi')\nhi")
    add_table(doc, [["输入", "输出"], ["1", "1"]])
    doc.save("report.docx")
    export_pdf("report.docx")          # only when a PDF is requested
"""

import shutil
import subprocess
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("python-docx not installed. Run: pip install python-docx")


def _run_cjk(run, cjk):
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), cjk)


def set_cjk(doc, style_name, latin, cjk, size):
    style = doc.styles[style_name]
    style.font.name = latin
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), cjk)


def new_report(body_latin="Times New Roman", body_cjk="宋体", body_size=12,
               head_latin="Arial", head_cjk="黑体"):
    """Return a Document with body and heading fonts set for CJK."""
    doc = Document()
    set_cjk(doc, "Normal", body_latin, body_cjk, body_size)
    for lvl in range(1, 4):
        name = f"Heading {lvl}"
        if name in doc.styles:
            set_cjk(doc, name, head_latin, head_cjk, body_size + 4 - lvl)
    return doc


def add_code(doc, text, mono="Consolas", cjk="宋体", size=10.5):
    """Add code or program output as a monospace paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = mono
    run.font.size = Pt(size)
    _run_cjk(run, cjk)
    return p


def add_table(doc, rows, header=True, style="Table Grid"):
    """Add a real grid table from a list of row lists."""
    if not rows:
        return None
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = style
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            run = table.cell(i, j).paragraphs[0].add_run(str(cell))
            _run_cjk(run, "宋体")
            if header and i == 0:
                run.bold = True
    return table


def add_picture(doc, path, width_inches=5.5):
    """Insert a real image file. Never invent images."""
    doc.add_picture(path, width=Inches(width_inches))


def export_pdf(docx_path, outdir="."):
    """Convert a .docx to .pdf with LibreOffice. Only when PDF is requested."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        sys.exit("LibreOffice not found; cannot export PDF.")
    subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                    "--outdir", outdir, docx_path], check=True)
